### CV Builder, the application will ask the user a series of questions either in text or speech and then build a cv for the user at the end.###

from tkinter.ttk import Style
from docx import Document
from docx.shared import Inches

#Openning the document
document = Document()

#Adding the profile picture and setting the size
document.add_picture(
    'johnny.png',
    width=Inches(2.0)
)

# Asking for basic details
name = input('What is your Name? ')
phone_number = input('What is your Phone Number? ')
email = input('What is your Email? ')

document.add_paragraph(
    f'{name} | {phone_number} | {email}'
)


# About me section 
print('-----------------------About me----------------------')
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself? ')
)

#Work experiences
print('------------------------Work Experience--------------')
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company: ')
from_date = input('From Date: ')
to_date = input('To Date: ')


p.add_run(f'{company} ').bold= True
p.add_run(f'{from_date} - {to_date} \n').italic = True

experience_details = input(
    f'Describe your experience at {company} '
)
p.add_run(f'{experience_details} \n \n')

#More experiences

while True:
    more_experience = input(
        'Do you have more experiences? (Yes or No) '
    )
    if more_experience.lower() == 'yes':
        company = input('Enter company: ')
        from_date = input('From Date: ')
        to_date = input('To Date: ')


        p.add_run(f'{company} ').bold= True
        p.add_run(f'{from_date} - {to_date} \n').italic = True

        experience_details = input(
            f'Describe your experince at {company} '
        )
        p.add_run(f'{experience_details} \n \n')
    else:
        break


#Skills Section
print('---------------------------------Skills----------------')
document.add_heading('Skills')
s = document.add_paragraph()
s.style = 'List Bullet'

skill = input(
    'Enter a skill: '
)

s.add_run(skill)

while True:
    more_skills = input(
        'Do you have more skills to add? (Yes/No) '
    )
    if more_skills.lower() == 'yes':
        
        

        skill = input(
            'Enter a skill: '
        )
        s = document.add_paragraph()
        s.style = 'List Bullet'
        s.add_run(skill)


    else:
        break







document.save('cv.docx')